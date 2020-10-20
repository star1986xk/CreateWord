# -*-coding: UTF-8 -*-
from datetime import datetime
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Inches, Pt, RGBColor
import openpyxl
from PyQt5.QtWidgets import QApplication, QFrame, QFileDialog, QMessageBox
from win import Ui_Form
from PyQt5.QtCore import Qt
import sys
import configparser


def font_style(run, size, name, color):
    # run.font.bold = True  # 黑体
    # run.font.italic = True  # 斜体
    # run.font.underline = True  # 下划线
    # run.font.strike = True  # 删除线
    # run.font.shadow = True  # 阴影
    run.font.size = Pt(size)  # 字体大小
    run.font.color.rgb = RGBColor(*color)  # 颜色
    run.font.name = name  # 字体
    r = run._element.rPr.rFonts  # 中文字体
    r.set(qn('w:eastAsia'), name)


def set_table_data(table, data, row, size, name, color):
    width_dict = {0: 6, 1: 2, 2: 1, 3: 1, 4: 2, 5: 1}
    tr = table.rows[row]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), "300")
    trHeight.set(qn('w:hRule'), "atLeast")
    trPr.append(trHeight)
    for n, v in enumerate(data):
        table.cell(row, n).paragraphs[0].text = v
        table.cell(row, n).width = Inches(width_dict.get(n))
        for run in table.cell(row, n).paragraphs[0].runs:
            font_style(run, size, name, color)


def get_excel(path):
    wb = openpyxl.load_workbook(path)
    ws = wb.active
    datas = []
    id = None
    date = None
    name = None
    for row in list(ws.rows)[1:]:
        if row[0].value:
            id = row[0].value
            date = row[1].value.strftime("%Y/%m/%d")
            name = row[9].value
            datas.append([
                str(row[3].value) if row[3].value else '',
                str(row[4].value) if row[4].value else '',
                str(row[5].value) if row[5].value else '',
                str(row[8].value) if row[8].value else '',
                str(row[7].value) if row[7].value else '',
                ''
            ])
        else:
            yield id, date, name, datas, True
            datas = []
            id = None
            date = None
            name = None
    yield id, date, name, datas, False


def create_word():
    doc = Document()
    # 设置页面大小
    section = doc.sections[0]
    section.page_width = Cm(21.5)
    section.page_height = Cm(13.9)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    return doc


def add_run(paragraph, text, size, name, color):
    run = paragraph.add_run(text)
    font_style(run, size, name, color)


def set_word(doc, id, name, date, head, business, mail, mailname, mail2, put, space, space2):
    # 第一行
    paragraph1 = doc.add_paragraph()
    paragraph1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_run(paragraph1, head, 12, '微软雅黑', (192, 0, 0))
    # 第二行
    paragraph2 = doc.add_paragraph()
    add_run(paragraph2,
            '业务：' + business + '     联系邮箱：' + mail + '     ' + mailname + '：' + mail2 + space + '单号：' + id,
            7.5, '微软雅黑', (0, 0, 0))
    # 第三行
    paragraph3 = doc.add_paragraph()
    add_run(paragraph3,
            '收货单位：',
            7.5, '微软雅黑', (0, 0, 0))
    add_run(paragraph3,
            put,
            7.5, '微软雅黑', (192, 0, 0))
    add_run(paragraph3,
            space2 + '制单员：' + name + '      制单日期：' + date,
            7.5, '微软雅黑', (0, 0, 0))


def set_word_end(doc, push, put2, space3, space4):
    data = ['注意：批量印刷前请细心对稿，如需重新制稿及制版请及时向我司提出反馈 批量印刷后所造成胡损失，我司概不负责。',
            '送货单位：' + push + space3 + '收货单位：' + put2,
            '经办人：' + space4 + '经办人：',
            ]
    for li in data:
        paragraph = doc.add_paragraph()
        add_run(paragraph,
                li,
                7.5, '微软雅黑', (0, 0, 0))


def create_table(doc, datas):
    # 表格
    table = doc.add_table(rows=6 if len(datas) < 6 else len(datas) + 1, cols=6, style='Table Grid')
    table.autofit = True
    data = ['产品名称（型号）', '物料编码', '物料规格', '单位', '总面积（平方厘米）', '备 注']
    set_table_data(table, data, 0, 7.5, '微软雅黑', (192, 0, 0))
    for n, data in enumerate(datas):
        set_table_data(table, data, n + 1, 7.5, '微软雅黑', (0, 0, 0))


class MainWindow(QFrame, Ui_Form):
    def __init__(self, parent=None):
        super(MainWindow, self).__init__(parent)

        self.setupUi(self)
        self.retranslateUi(self)

        self.setWindowFlags(
            Qt.Window | Qt.FramelessWindowHint | Qt.WindowSystemMenuHint | Qt.WindowMinimizeButtonHint | Qt.WindowMaximizeButtonHint)  # 窗口无边框

        self.pushButton_4.clicked.connect(self.close)
        self.pushButton.clicked.connect(self.open_excel)
        self.pushButton_2.clicked.connect(self.run)

        self.get_config()

    def get_config(self):
        cf = configparser.RawConfigParser()
        cf.read("./config.ini", encoding='utf-8-sig')
        self.head = cf.get('settings', '表头')
        self.business = cf.get('settings', '业务')
        self.mail = cf.get('settings', '联系邮箱')
        self.mailname = cf.get('settings', '邮箱名称')
        self.mail2 = cf.get('settings', '联系邮箱2')
        self.put = cf.get('settings', '收货单位')
        self.push = cf.get('settings', '送货单位')
        self.put2 = cf.get('settings', '收货单位2')
        self.space = ' ' * int(cf.get('settings', '空格'))
        self.space2 = ' ' * int(cf.get('settings', '空格2'))
        self.space3 = ' ' * int(cf.get('settings', '空格3'))
        self.space4 = ' ' * int(cf.get('settings', '空格4'))

    def run(self):
        path = self.lineEdit.text()
        if path:
            doc = create_word()
            for id, date, name, datas, flag in get_excel(path):
                if id:
                    set_word(doc, id, name, date, self.head, self.business, self.mail, self.mailname, self.mail2,
                             self.put, self.space, self.space2)
                    create_table(doc, datas)
                    set_word_end(doc, self.push, self.put2, self.space3, self.space4)
                if flag:
                    doc.add_page_break()
            file_name = datetime.now().strftime('%Y-%m-%d %H,%M,%S') + '.docx'
            doc.save(file_name)
            QMessageBox.information(None, "提示", "生成 " + file_name + " 完成！")

    def open_excel(self):
        self.file_path_excel, _ = QFileDialog.getOpenFileName(self, '选择文件', '', 'Excel files(*.xlsx;)')
        self.lineEdit.setText(self.file_path_excel)

    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self.m_flag = True
            self.m_Position = event.globalPos() - self.pos()  # 获取鼠标相对窗口的位置
            event.accept()

    def mouseMoveEvent(self, QMouseEvent):
        if Qt.LeftButton and self.m_flag:
            self.move(QMouseEvent.globalPos() - self.m_Position)  # 更改窗口位置
            QMouseEvent.accept()

    def mouseReleaseEvent(self, QMouseEvent):
        self.m_flag = False


if __name__ == '__main__':
    app = QApplication(sys.argv)
    win = MainWindow()
    win.show()
    sys.exit(app.exec_())
